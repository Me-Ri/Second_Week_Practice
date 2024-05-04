from docxtpl import DocxTemplate, R
from docx import Document
import os, sys
from docx.shared import Pt

#############################################################################

def delete_file(name_file):
    if(os.path.exists(name_file)):
	    os.remove(name_file)

#############################################################################

delete_file("doc-final.docx")

php_mas = sys.argv

php_mas.pop(0)
php_mas[6] = php_mas[6].split(",")

veriables = [   
	        "initial_manager_OPOP",
            "initial_manager_practice",
            "code_direction",
            "name_direction",
            "course",
            "group",
			"",
            "comment",
            "number_order",
            "date_order",
            "date_practice",
			"view",
			"type",
			"count_good",
			"count_bad"
			
		    ]

mark = []
good_student = []
bad_student = []
badreason = []

def good_or_bad_student():
    count_bad = 0
    count_good = 0
    for i in range(0, len(php_mas[6])):
        mas = php_mas[6][i].split("@")
        mark.append(mas[0])

        if(int(mas[0]) < 3):
            badreason.append(mas[1])
            bad_student.append(mas[2])
            count_bad += 1
        else:
            good_student.append(mas[2])
            count_good += 1
    php_mas.append(count_good)
    php_mas.append(count_bad)
			
def initials():
	for i in range(0, 2):
		mas = php_mas[i].split()
		init = mas[0] + " " + mas[1][0] + "." + mas[2][0] + "."
		php_mas[i] = init

#############################################################################

def fill_table():
	doc = Document("образец документа 2.docx")

	good_or_bad_student()

	doc_table = doc.tables[0]
	
	for i in range(0, len(good_student)):
		new_row = doc_table.add_row()
		new_row.cells[0].text = " " + str(i+ 1)
		new_row.cells[1].text = " " + good_student[i]
		new_row.cells[2].text = " " + "ЮГУ"
		new_row.cells[3].text = " " + "г.Ханты-Мансийск"
		new_row.cells[4].text = "-"
		new_row.cells[5].text = "-"
		new_row.cells[6].text = "      " + php_mas[1]

	doc_table = doc.tables[1]

	for i in range(0, len(bad_student)):
		new_row = doc_table.add_row()
		new_row.cells[0].text = " " + str(i + 1)
		new_row.cells[1].text = "     " + bad_student[i]
		new_row.cells[2].text = badreason[i]

	doc.save("table.docx")

#############################################################################



initials()
fill_table()

doc = DocxTemplate("table.docx")



contex = {veriables[i]:php_mas[i] for i in range(len(veriables))}

doc.render(contex)

doc.save("doc-final.docx")

delete_file("table.docx")
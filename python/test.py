from docxtpl import DocxTemplate, R
from docx import Document
import os, sys
import csv

#############################################################################

def delete_file(name_file):
    if(os.path.exists(name_file)):
	    os.remove(name_file)

#############################################################################

delete_file("doc-final.docx")

php_mas = [
            sys.argv[1],
            sys.argv[2],
            sys.argv[3],
            sys.argv[4],
            sys.argv[5],
            sys.argv[6],
            sys.argv[7],
            [],
            sys.argv[8],
            sys.argv[9],
            sys.argv[10],
            sys.argv[11],
            sys.argv[12],
            sys.argv[13],
            sys.argv[14],
            sys.argv[15],
            sys.argv[16],
            sys.argv[17].lower(),
            sys.argv[18].lower(),
            sys.argv[19].lower(),
            sys.argv[20],
            sys.argv[21]
            
]


veriables = [   
	        "initials_manager_enterprise",
            "initials_manager_organization",
            "initials_manager_ugu",
            "initials_student",
            "FNP_student",
            "university",
            "course",
            "group",
            "task",
            "place_practice",
            "post_manager_enterprise",
            "type_practice",
            "dates_practice",
            "post_manager_organization",
            "direction_training",
            "start_dates",
            "start_month",
            "finish_dates",
            "finish_month",
            "post_manager_ugu",
            #"post_manager_organization",
            "name_practice",
            "address_orginization",
            "quality",
            "problems",
            "completing_task",
            "comments",
            "grade",
            
		    ]

############################################################################


def fill_table():
    mas_issues = []

    with open('issues.csv', 'r',  encoding='utf-8-sig', newline='') as csvfile:
        spamreader = csv.reader(csvfile, delimiter=',', quotechar='|')
        for row in spamreader:
            mas_issues.append([row[9], row[13], row[7]])

    doc_table = Document("образец документа.docx")

    doc_table.tables[4]

    name_sername_mas = php_mas[3].split()
    name_sername = name_sername_mas[1] + " " + name_sername_mas[0]

    count_row = 3
    for i in range(3, len(mas_issues) + 2):
        if(mas_issues[i - 2][0] == name_sername):
            for j in range(0, 2):
                if (j == 0):
                    doc_table.tables[4].cell(count_row, j).text = mas_issues[i - 2][1].replace("/", ".")
                if (j == 1):
                    doc_table.tables[4].cell(count_row, j).text = mas_issues[i - 2][2]
                    php_mas[7].append(mas_issues[i - 2][2])
            count_row += 1

    doc_table.save("table.docx")

############################################################################

fill_table()
doc = DocxTemplate("table.docx")

def generate_tasks():
    for i in range (0, len(php_mas[8])):
        ok = str(i + 1) + ') ' + php_mas[8][i]
        yield ok
		
def initials():
	php_mas.insert(4, php_mas[3])
	for i in range(0, 4):
		mas = php_mas[i].split()
		init = mas[0] + " " + mas[1][0] + "." + mas[2][0] + "."
		php_mas[i] = init

def dates():
      mas_month = ["январь", "февраль", "март", "апрель", "май", "июнь", "июль", "август", "сентябрь", "октябрь", "ноябрь", "декабрь"]
      php_mas.insert(15, php_mas[12][0:2]) 
      php_mas.insert(16, mas_month[int(php_mas[12][3:5]) - 1])
      php_mas.insert(17, php_mas[12][11:13])
      php_mas.insert(18, mas_month[int(php_mas[12][14:16]) - 1])

initials()
dates()

contex = {veriables[i]:php_mas[i] for i in range(len(veriables))}
contex['task'] = R("\n\t".join(generate_tasks()))
# contex = {"group":sys.argv}

doc.render(contex)

doc_puti = "doc-final_" + sys.argv[22] + ".docx"

doc.save(doc_puti)

delete_file("table.docx")